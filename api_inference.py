from langchain_community.llms import VertexAI
from langchain_community.embeddings import VertexAIEmbeddings
from google.cloud import storage
import docx
import langchain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import ConversationalRetrievalChain
from datetime import datetime
from langchain_google_vertexai import VertexAI
from langchain.text_splitter import CharacterTextSplitter
from langchain.agents.agent_types import AgentType
from langchain_community.vectorstores import FAISS
from PyPDF2 import PdfReader
from typing import List
import json
import pandas as pd
import os
from google.cloud import storage
from google.cloud import bigquery
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from google.cloud.bigquery import Client
from tqdm import tqdm
import pandas as pd
from typing import Any, List, Dict, Union
from langchain.docstore.document import Document
from langchain.schema import Document
from langchain.vectorstores.base import VectorStoreRetriever
from google.cloud import bigquery
import io
from langchain_core.pydantic_v1 import BaseModel, Extra, Field, root_validator
from langchain.prompts.example_selector import SemanticSimilarityExampleSelector
from langchain_community.vectorstores import Chroma
from datetime import datetime
from langchain.prompts import PromptTemplate, FewShotPromptTemplate
from langchain.memory import ConversationBufferMemory
from adapt_api import *
from helper import *
from function_declarations import agent


llm=VertexAI(model_name='gemini-pro', temperature = 0.5)
embeddings = VertexAIEmbeddings(model_name="textembedding-gecko@003") 

HISTORY_ACTION = "NO_ACTION" #based on this global variable the chat history will either be cleared on kept untouched
CHAT_HISTORY = [] #To store intermediate chat history for api calls or clear

class QuestionClassifier:
    
    def __init__(self):
        self.file_name = 'api_class_exmpls.json'
        self.llm=VertexAI(model_name="text-bison")
        self.embeddings = VertexAIEmbeddings(model_name="textembedding-gecko@latest")
        self.examples = self.get_examples()
        
    def append_examples(self,example: dict) -> list[dict]:
        """This function adds additional example prompts as well as their classificatin
    Parameters:
    example (dict): new example dict in the format of {"input": , "output": }
    Returns:
    list[dict]: examples in the form of list of dicts
    """

        # Read the existing data
        with open(self.file_name, 'r') as file:
            examples = json.load(file)

        # Append the new dictionary to the list
        examples.append(example)

        # Write the updated list back to the file
        with open(self.file_name, 'w') as file:
            json.dump(examples, file)

        print("Examples updated")

        return examples

    def get_examples(self) -> list[dict]:
        """
        This function returns example prompts
    Returns:
    list[dict]: examples in the form of list of dicts
        """
        with open(self.file_name, 'r') as file:
            examples = json.load(file)
        return examples


    def qc_chain(self, question:str) -> str:
        """This function classifies the question asked by user
    Parameters:
    question (str): question from user
    Returns:
    str: category of question
    """
        
        example_prompt = PromptTemplate(
        input_variables=["question", "output"],
        template="question: {question}\noutput: {output}",)

        example_selector = SemanticSimilarityExampleSelector.from_examples(
        # The list of examples available to select from.
        self.examples,
        # The embedding class used to produce embeddings which are used to measure semantic similarity.
        self.embeddings,
        # The VectorStore class that is used to store the embeddings and do a similarity search over.
        FAISS,
        # The number of examples to produce.
        k=5,
        )
        similar_prompt = FewShotPromptTemplate(
        # We provide an ExampleSelector instead of examples.
        example_selector=example_selector,
        example_prompt=example_prompt,
        prefix="Given the input question and examples please classify the question and give the output",
        suffix=" Classify the input below based on above examples give the output as result \n rule: only give output never explain \n rule: the input can be in the form of question. So, just classify.  \n question: {question}\n output:",
        input_variables=["question"],
        )
        classification_chain = similar_prompt | self.llm #instantiating the chain
        question_category = classification_chain.invoke({"question":question}).strip()
        return question_category
    
qclass = QuestionClassifier()

class Document:
    def __init__(self, page_content, metadata):
        self.page_content = page_content
        self.metadata = metadata
    def __str__(self):
        return f"Document(page_content='{self.page_content}', metadata={self.metadata})"
    def __repr__(self):
        return self.__str__()


class FilteredRetriever(VectorStoreRetriever):
    vectorstore: VectorStoreRetriever
    search_type: str = "similarity"
    search_kwargs: dict = Field(default_factory=dict)
    def get_relevant_documents(self, query: str) -> List[Document]:
        global function_call
        function_call = 'dummy'#qclass.qc_chain(query)
        docs = []
        """print('query: ', query)
        print('function_call: ',function_call)
        docs = []
        if function_call in globals():
            global HISTORY_ACTION
            response=globals()[function_call](query)
            output,HISTORY_ACTION  = response
            from global_definitions import update_global
            if '{' in output:
                API_RESPONSE = output
                update_global({'API_RESPONSE' : API_RESPONSE})
                docs.append(Document(page_content = 'Do not return anything', metadata= {}))
            else:
            
                docs.append(Document(page_content = output, metadata= {}))
        else:
            print("Wrong classification !!")"""
        response = agent(query)
        docs.append(Document(page_content = response, metadata= {}))

        return docs
    
vector_db = FAISS.load_local('docs_gecko_store',embeddings, allow_dangerous_deserialization = True)
filtered_retriever = FilteredRetriever(vectorstore = vector_db.as_retriever())

q_template = """Examine the user's query and the provided context. Decide if the content of context should be relayed directly, summarized, or if it contains a question that should be posed to the user.
 
Conversation Context:
- Chat History: {chat_history}
- User Query: {question}
- Context: {context} (The context might contain direct information, require summarization, or include a question to be asked to the user.)
 
Task:
-If the context directly answers the user's query, present that information. If the context is broad or contains detailed information, summarize the key points. If the context contains a question that is relevant to the conversation or necessary for further clarification, present that question to the user.
Rules:
- If the context has a question directly use that question as a response. if context aska for confirmation give the context directly as output.
- Never justfy or explain your reason for displaying an output.
- If the context does not have question then summarize the context based on the users query and provide steps if necessary in detailed way.

 
Output:
- [Depending on the context content and the rules provide a direct response or pose a relevant question to the user ]
- [if you cannot do the above task just return the context]
"""
PROMPT=PromptTemplate(input_variables=["context","question","chat_history"],template=q_template)
memory = ConversationBufferMemory(memory_key='chat_history',return_docs=False, return_messages=True, input_key='question', output_key='answer')

def get_new_prompt():
    custom_template = """Given the user's latest input and the preceding chat history, rephrase the user's input into a standalone, clear, and contextually appropriate action or question.
 
- Chat History: {chat_history}
- Latest Human Input: {question}
 
Rules:
- Rephrase the user's latest input to make it a clear, standalone statement or question that is relevant to the ongoing conversation context.
- Only add 'Yes' in the beginning of the output if and only if the latest Human input has 'Yes' in it.

Output:[Rephrase the Latest Human Input to be contextually appropriate and standalone]
if you cannot do the above task just return the latest human input
"""
    CUSTOM_QUESTION_PROMPT = PromptTemplate.from_template(custom_template)
    #print(CUSTOM_QUESTION_PROMPT)
    return CUSTOM_QUESTION_PROMPT

chain = ConversationalRetrievalChain.from_llm(
    llm=llm,
    retriever=filtered_retriever,
    memory = memory,
    verbose=True,
    return_source_documents=True,
    condense_question_prompt=get_new_prompt(),
    combine_docs_chain_kwargs={"prompt": PROMPT})
 
def api_chain(user_input):
    global HISTORY_ACTION
    global function_call
    chat_history,question = user_input #user input is a tuple of chat history and query
    question = extract_identifier(question) # extract dtn,ordernumber or any alphanumeric
    print(question)
    if 'yes' not in question.lower() and 'sure' not in question.lower() and 'confirm' not in question.lower():
        CHAT_HISTORY.clear() 
    res = chain({"question": question, "chat_history":CHAT_HISTORY})
    answer = res["answer"]
    print(answer)
    answer = answer.replace('Pose a question to the user: ','').replace('Question to the user: ','').replace('Pose the following question to the user: ','').replace("- Assistant's Response: ",'').replace('Question: ','')
    if HISTORY_ACTION == 'CLEAR_HISTORY': # if history action is clear history the list should be cleared and the variable should be reset
        HISTORY_ACTION = "NO_ACTION" #resetting history action
        CHAT_HISTORY.clear() 
    else:
        CHAT_HISTORY.append({'user':question, 'model':answer}) #appending chat history
        
    from global_definitions import API_RESPONSE 
    #Check if the response data is in JSON format and format it accordingly
    if API_RESPONSE != 'initial':
        print('api response', API_RESPONSE)
        try:
            response_data = API_RESPONSE.replace("'", '"')
            # Format response as JSON
            formatted_response = {
                "Content": [function_call],
                "responsedata": [json.loads(response_data)]
            }
            print('formatted_response', formatted_response)
            return json.dumps(formatted_response)
        except json.JSONDecodeError:
            return answer
    else:
        return answer
    #return answer,function_call
